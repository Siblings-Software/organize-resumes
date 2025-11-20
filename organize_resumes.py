#!/usr/bin/env python3
"""
Resume Organizer Script
Automatically moves resume files from root folder to appropriate role folders
based on content analysis of technologies and roles mentioned in the resumes.
"""

import os
import re
import shutil
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from collections import defaultdict, Counter
import io


class ResumeOrganizer:
    def __init__(self, root_dir: str, dry_run: bool = True, manual_mappings: Optional[dict] = None):
        """
        Initialize the Resume Organizer.
        
        Args:
            root_dir: Root directory containing CVs
            dry_run: If True, only show what would be done without actually moving files
            manual_mappings: Dictionary mapping filenames to target folders
        """
        self.root_dir = Path(root_dir)
        self.dry_run = dry_run
        self.manual_mappings = manual_mappings or {}
        self.resume_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        self.role_folders = []
        self.moved_files = []
        self.matched_files = []  # Track matched files even in dry-run
        self.unmatched_files = []
        self.created_folders = []
        
        # Build technology and role keyword mapping
        self.role_keywords = self._build_role_keywords()
        
    def _build_role_keywords(self) -> Dict[str, List[str]]:
        """
        Build a mapping of folder names to keywords that indicate that role/technology.
        Based on existing folder structure and common technologies.
        """
        return {
            'Python': [
                'python', 'django', 'flask', 'fastapi', 'pandas', 'numpy', 'scikit-learn',
                'tensorflow', 'pytorch', 'keras', 'celery', 'sqlalchemy', 'pytest'
            ],
            'JavaScript': [
                'javascript', 'js', 'node.js', 'nodejs', 'react', 'vue', 'angular',
                'typescript', 'express', 'next.js', 'nuxt', 'jquery', 'webpack', 'babel'
            ],
            'React': [
                'react', 'react.js', 'reactjs', 'redux', 'next.js', 'gatsby', 'jsx'
            ],
            'Node': [
                'node.js', 'nodejs', 'express', 'nestjs', 'koa', 'socket.io', 'mongo'
            ],
            'Java': [
                'java', 'spring', 'spring boot', 'hibernate', 'jpa', 'maven', 'gradle',
                'jsp', 'servlet', 'javaee', 'j2ee', 'struts'
            ],
            '.NET': [
                '.net', 'dotnet', 'c#', 'csharp', 'asp.net', 'aspnet', 'entity framework',
                'blazor', 'xamarin', 'wcf', 'linq', 'nuget'
            ],
            'DevOPS': [
                'devops', 'docker', 'kubernetes', 'k8s', 'jenkins', 'ci/cd', 'ci cd',
                'terraform', 'ansible', 'chef', 'puppet', 'aws', 'azure', 'gcp',
                'github actions', 'gitlab ci', 'prometheus', 'grafana', 'elasticsearch',
                'kibana', 'nginx', 'apache', 'linux', 'bash', 'shell scripting'
            ],
            'GoLang': [
                'go', 'golang', 'goroutine', 'gin', 'echo', 'beego'
            ],
            'PHP': [
                'php', 'laravel', 'symfony', 'codeigniter', 'wordpress', 'drupal', 'magento'
            ],
            'WordPress': [
                'wordpress', 'wp', 'woocommerce', 'php', 'theme development', 'plugin development'
            ],
            'Magento': [
                'magento', 'magento 2', 'ecommerce', 'php'
            ],
            'Mobile': [
                'ios', 'swift', 'objective-c', 'android', 'kotlin', 'react native',
                'flutter', 'xamarin', 'mobile app', 'app development', 'ios engineer',
                'ios developer', 'android developer', 'mobile developer'
            ],
            'Kotlin': [
                'kotlin', 'android', 'coroutines', 'ktor'
            ],
            'Unity': [
                'unity', 'unity3d', 'c#', 'game development', 'unity engine'
            ],
            'QA': [
                'qa', 'quality assurance', 'testing', 'selenium', 'cypress', 'jest',
                'junit', 'test automation', 'manual testing', 'automated testing',
                'test cases', 'bug tracking', 'jira', 'testrail'
            ],
            'Fullstack': [
                'full stack', 'fullstack', 'full-stack', 'mern', 'mean', 'mevn',
                'react', 'node', 'mongodb', 'postgresql', 'mysql'
            ],
            'Blockchain': [
                'blockchain', 'ethereum', 'solidity', 'web3', 'smart contracts',
                'bitcoin', 'cryptocurrency', 'defi', 'nft'
            ],
            'Data Engineer': [
                'data engineer', 'etl', 'data pipeline', 'apache spark', 'hadoop',
                'airflow', 'data warehouse', 'sql', 'big data', 'kafka', 'flink'
            ],
            'Data Security': [
                'security', 'cybersecurity', 'penetration testing', 'vulnerability assessment',
                'owasp', 'ssl/tls', 'encryption', 'security audit', 'compliance'
            ],
            'WebOps': [
                'webops', 'web operations', 'deployment', 'hosting', 'cdn', 'load balancing'
            ],
            'Salesforce': [
                'salesforce', 'apex', 'soql', 'lightning', 'salesforce admin',
                'salesforce developer', 'sfdc'
            ],
            'SAP': [
                'sap', 'abap', 'sap erp', 'sap fico', 'sap mm', 'sap sd'
            ],
            'Dynamics': [
                'dynamics', 'dynamics 365', 'dynamics crm', 'power platform', 'powerapps'
            ],
            'Contentful': [
                'contentful', 'cms', 'headless cms', 'content management'
            ],
            'Clojure': [
                'clojure', 'clojurescript'
            ],
            'Delphi': [
                'delphi', 'pascal', 'firemonkey'
            ],
            'Ruby on Rails': [
                'ruby', 'rails', 'ruby on rails', 'ror', 'sinatra'
            ],
            'Project Manager': [
                'project manager', 'pmp', 'agile', 'scrum master', 'project management',
                'stakeholder', 'project planning', 'risk management'
            ],
            'Scrum Master': [
                'scrum master', 'agile coach', 'sprint', 'scrum', 'kanban', 'sprint planning'
            ],
            'Product Manager': [
                'product manager', 'product management', 'product owner', 'roadmap',
                'user stories', 'product strategy'
            ],
            'Technical Leader': [
                'technical lead', 'tech lead', 'team lead', 'engineering manager',
                'architecture', 'system design', 'leadership'
            ],
            'Analista Funcional': [
                'analista funcional', 'business analyst', 'functional analyst',
                'requirements analysis', 'bpmn'
            ],
            'UX UI': [
                'ux', 'ui', 'user experience', 'user interface', 'design', 'figma',
                'adobe xd', 'sketch', 'prototyping', 'wireframing'
            ],
            'Director de Arte': [
                'art director', 'director de arte', 'creative director', 'graphic design',
                'art direction'
            ],
            'Marketing': [
                'marketing', 'digital marketing', 'seo', 'sem', 'google ads', 'facebook ads',
                'content marketing', 'social media', 'analytics'
            ],
            'Sales Engineer': [
                'sales engineer', 'pre-sales', 'technical sales', 'solution architect'
            ],
            'Customer Experience': [
                'customer experience', 'cx', 'customer success', 'customer support'
            ],
            'Technical Writer': [
                'technical writer', 'technical writing', 'documentation', 'api documentation'
            ],
            'Recruiters': [
                'recruiter', 'talent acquisition', 'hr', 'human resources', 'recruiting'
            ],
            'Logística': [
                'logistics', 'logística', 'supply chain', 'inventory management'
            ],
            'Pasantes': [
                'intern', 'internship', 'pasante', 'trainee'
            ],
            'Trainees': [
                'trainee', 'training', 'junior', 'entry level'
            ],
            'Juniors': [
                'junior', 'entry level', 'graduate', 'recent graduate'
            ],
            'Piscólogas': [
                'psychology', 'psicología', 'psychologist', 'mental health'
            ],
            'Fuera de la Industria': [
                'outside tech', 'non-tech', 'other industry'
            ]
        }
    
    def get_role_folders(self) -> List[Path]:
        """Get all role folder directories."""
        if not self.role_folders:
            for item in self.root_dir.iterdir():
                if item.is_dir() and not item.name.startswith('.'):
                    excluded = {'desktop.ini', '__pycache__', '.git'}
                    if item.name not in excluded:
                        self.role_folders.append(item)
        return self.role_folders
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            # Try PyPDF2 first
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                # Try pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            text += page.extract_text() + "\n" if page.extract_text() else ""
                    return text
                except ImportError:
                    print(f"[WARNING] No PDF library found. Install PyPDF2 or pdfplumber: pip install PyPDF2")
                    return ""
        except Exception as e:
            print(f"[WARNING] Error reading PDF {file_path.name}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            import docx
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            return text
        except ImportError:
            print(f"[WARNING] python-docx not found. Install it: pip install python-docx")
            return ""
        except Exception as e:
            print(f"[WARNING] Error reading DOCX {file_path.name}: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from resume file based on extension."""
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except Exception as e:
                print(f"[WARNING] Error reading TXT {file_path.name}: {e}")
                return ""
        elif ext == '.doc':
            # Old .doc format - harder to read without special libraries
            print(f"[INFO] .doc format not fully supported for {file_path.name}, trying as text")
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                    # Try to extract readable text (basic approach)
                    text = content.decode('utf-8', errors='ignore')
                    # Remove non-printable characters
                    text = ''.join(c for c in text if c.isprintable() or c.isspace())
                    return text
            except Exception as e:
                print(f"[WARNING] Error reading DOC {file_path.name}: {e}")
                return ""
        
        return ""
    
    def analyze_resume_content(self, text: str) -> Dict[str, int]:
        """
        Analyze resume text and count keyword occurrences for each role/technology.
        
        Returns:
            Dictionary mapping role names to their keyword count scores
        """
        text_lower = text.lower()
        scores = defaultdict(int)
        
        for role, keywords in self.role_keywords.items():
            for keyword in keywords:
                # Count occurrences of keyword (as whole word)
                pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
                matches = len(re.findall(pattern, text_lower))
                scores[role] += matches
        
        return dict(scores)
    
    def find_target_folder_from_content(self, file_path: Path) -> Optional[Tuple[Path, str, int]]:
        """
        Analyze resume content to determine the best matching folder.
        Also checks filename for keywords if content extraction fails.
        
        Returns:
            Tuple of (target_folder_path, role_name, score) or None if no match
        """
        # Try to extract text from file
        text = self.extract_text_from_file(file_path)
        
        # If content extraction failed or returned little text, also analyze filename
        filename_text = ""
        if not text or len(text.strip()) < 50:
            filename_text = file_path.stem.lower() + " " + file_path.name.lower()
            if text:
                text = text + " " + filename_text
            else:
                text = filename_text
        
        # Analyze content (including filename if needed)
        scores = self.analyze_resume_content(text)
        
        if not scores:
            return None
        
        # Find the role with highest score
        best_role = max(scores.items(), key=lambda x: x[1])
        role_name, score = best_role
        
        # Lower threshold if we only matched from filename
        min_score = 1 if filename_text else 2
        
        # Only proceed if score is significant
        if score < min_score:
            return None
        
        # Check if folder exists
        target_folder = self.root_dir / role_name
        
        # If folder doesn't exist, we'll create it
        if not target_folder.exists():
            return (target_folder, role_name, score)
        
        # Folder exists
        return (target_folder, role_name, score)
    
    def create_folder_if_needed(self, folder_path: Path) -> bool:
        """Create folder if it doesn't exist."""
        if not folder_path.exists():
            if not self.dry_run:
                try:
                    folder_path.mkdir(parents=True, exist_ok=True)
                    self.created_folders.append(folder_path.name)
                    print(f"[CREATED] New folder: {folder_path.name}/")
                    return True
                except Exception as e:
                    print(f"[ERROR] Could not create folder {folder_path.name}: {e}")
                    return False
            else:
                self.created_folders.append(folder_path.name)
                print(f"[DRY RUN] Would create folder: {folder_path.name}/")
                return True
        return False
    
    def find_target_folder(self, file_path: Path) -> Optional[Tuple[Path, str, int]]:
        """
        Find the target role folder for a resume file.
        First checks manual mappings, then analyzes content.
        
        Returns:
            Tuple of (target_folder_path, role_name, score) or None
        """
        filename = file_path.name
        
        # Check manual mappings first
        if filename in self.manual_mappings:
            target_folder_name = self.manual_mappings[filename]
            target_folder = self.root_dir / target_folder_name
            self.create_folder_if_needed(target_folder)
            return (target_folder, target_folder_name, 999)  # High score for manual mappings
        
        # Analyze file content
        return self.find_target_folder_from_content(file_path)
    
    def should_move_file(self, file_path: Path) -> bool:
        """Check if a file should be moved (is a resume file)."""
        excluded_files = {
            'organize_resumes.py',
            'requirements.txt',
            'desktop.ini'
        }
        return (
            file_path.is_file() and
            file_path.suffix.lower() in self.resume_extensions and
            not file_path.name.startswith('.') and
            file_path.name not in excluded_files
        )
    
    def organize_resumes(self):
        """Main method to organize all resumes."""
        print(f"{'='*60}")
        print(f"Resume Organizer - Content Analysis Mode")
        print(f"{'DRY RUN MODE' if self.dry_run else 'LIVE MODE'}")
        print(f"{'='*60}\n")
        
        # Get all resume files in root
        root_files = [f for f in self.root_dir.iterdir() if self.should_move_file(f)]
        
        print(f"Found {len(root_files)} resume files in root folder\n")
        
        for file_path in root_files:
            result = self.find_target_folder(file_path)
            
            if result:
                target_folder, role_name, score = result
                
                # Ensure folder exists
                self.create_folder_if_needed(target_folder)
                
                destination = target_folder / file_path.name
                
                # Handle duplicate filenames
                counter = 1
                while destination.exists():
                    stem = file_path.stem
                    suffix = file_path.suffix
                    new_name = f"{stem} ({counter}){suffix}"
                    destination = target_folder / new_name
                    counter += 1
                
                self.matched_files.append((file_path.name, role_name, score))
                
                if self.dry_run:
                    print(f"[DRY RUN] Would move: {file_path.name}")
                    print(f"          -> {role_name}/ (score: {score})")
                    print()
                else:
                    try:
                        shutil.move(str(file_path), str(destination))
                        print(f"[MOVED] {file_path.name} -> {role_name}/ (score: {score})")
                        self.moved_files.append((file_path.name, role_name, score))
                    except Exception as e:
                        print(f"[ERROR] Error moving {file_path.name}: {e}")
            else:
                print(f"[UNMATCHED] No match found: {file_path.name}")
                self.unmatched_files.append(file_path.name)
        
        # Summary
        print(f"\n{'='*60}")
        print("SUMMARY")
        print(f"{'='*60}")
        print(f"Files matched: {len(self.matched_files)}")
        print(f"Files unmatched: {len(self.unmatched_files)}")
        print(f"Folders created: {len(self.created_folders)}")
        
        if self.created_folders:
            print(f"\nNew folders created:")
            for folder in self.created_folders:
                print(f"  - {folder}/")
        
        if self.unmatched_files:
            print(f"\nUnmatched files (may need manual placement):")
            for filename in self.unmatched_files[:20]:  # Show first 20
                print(f"  - {filename}")
            if len(self.unmatched_files) > 20:
                print(f"  ... and {len(self.unmatched_files) - 20} more")


def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Organize resume files into role folders based on content analysis.'
    )
    parser.add_argument(
        '--root-dir',
        type=str,
        default='.',
        help='Root directory containing CVs (default: current directory)'
    )
    parser.add_argument(
        '--execute',
        action='store_true',
        help='Actually move files (default: dry-run mode)'
    )
    parser.add_argument(
        '--mappings',
        type=str,
        help='Path to JSON file with manual filename-to-folder mappings'
    )
    
    args = parser.parse_args()
    
    # Load manual mappings if provided
    manual_mappings = {}
    if args.mappings:
        import json
        try:
            with open(args.mappings, 'r', encoding='utf-8') as f:
                manual_mappings = json.load(f)
            print(f"Loaded {len(manual_mappings)} manual mappings from {args.mappings}\n")
        except Exception as e:
            print(f"Warning: Could not load mappings file: {e}\n")
    
    # Add default mappings
    default_mappings = {
        "David Cano.pdf": "DevOPS",
        "Francisco Lala - Resume.pdf": "JavaScript"
    }
    manual_mappings.update(default_mappings)
    
    organizer = ResumeOrganizer(args.root_dir, dry_run=not args.execute, manual_mappings=manual_mappings)
    organizer.organize_resumes()
    
    if organizer.dry_run and not args.execute:
        print("\n[INFO] This was a dry run. Use --execute to actually move files.")
        print("[INFO] Install required libraries: pip install PyPDF2 python-docx (or pdfplumber)")


if __name__ == '__main__':
    main()
